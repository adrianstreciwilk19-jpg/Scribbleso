import html
import json
import os
import re
import sys
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, ttk


DEFAULT_FONT_SIZE = 11
DESCRIPTION_FONT_SIZE = 9
UI_FONT_SIZE = 13
UI_SMALL_FONT_SIZE = 11
SETTINGS_FILE = (
    Path(os.environ.get("APPDATA") or Path.home()) / "Scribbleso" / "settings.json"
)
ANONYMIZE_PATTERN = re.compile(
    r"""
    (?P<email>\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b)
    |
    (?P<address>
        (?<!\w)
        (?:
            (?i:ul\.?|ulica|al\.?|aleja|pl\.?|plac|os\.?|osiedle|rondo|rynek)\s+
            [A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż0-9.'-]+
            (?:\s+[A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż0-9.'-]+){0,5}
            \s+\d+[A-Za-z]?(?:[/-]\d+[A-Za-z]?)?
            (?:\s*,?\s*\d{2}-\d{3}\s+
                [A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+
                (?:\s+[A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+){0,3}
            )?
            |
            [A-ZĄĆĘŁŃÓŚŹŻ][A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+
            (?:\s+[A-ZĄĆĘŁŃÓŚŹŻ][A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+){0,4}
            \s+\d+[A-Za-z]?(?:[/-]\d+[A-Za-z]?)?
            \s*,?\s*\d{2}-\d{3}\s+
            [A-ZĄĆĘŁŃÓŚŹŻ][A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+
            (?:\s+[A-ZĄĆĘŁŃÓŚŹŻ][A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż.'-]+){0,3}
        )
        (?!\w)
    )
    |
    (?P<number_sequence>(?<!\w)\+?\d(?:[\s.-]?\d){6,}(?!\w))
    |
    (?P<identifier>
        (?<!\w)
        (?:
            (?=[A-Za-z0-9._/-]*\d)(?=[A-Za-z0-9._/-]*[._/-])
            [A-Za-z0-9]+(?:[._/-][A-Za-z0-9]+)+
            |
            (?=(?:[A-Za-z0-9]*\d){3,})(?=[A-Za-z0-9]*[A-Za-z])
            [A-Za-z0-9]{4,}
            |
            \d{3,}
        )
        (?!\w)
    )
    |
    (?P<person>
        (?<!\w)
        [A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,}(?:-[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})?
        \s+
        [A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,}(?:-[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})?
        (?:\s+[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,}(?:-[A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]{2,})?)?
        (?!\w)
    )
    """,
    re.VERBOSE,
)
PROTECTED_REFERENCE_PATTERN = re.compile(
    r"""
    (?<!\w)
    (?:
        (?i:(?:RZ|Z|R)-\d+-\d{2,4})
        |
        (?i:ZU\d+)
    )
    (?!\w)
    """,
    re.VERBOSE,
)
MONTH_NAME_TO_NUMBER = {
    "sty": 1,
    "styczeń": 1,
    "stycznia": 1,
    "jan": 1,
    "january": 1,
    "lut": 2,
    "luty": 2,
    "lutego": 2,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "marzec": 3,
    "marca": 3,
    "march": 3,
    "kwi": 4,
    "kwiecień": 4,
    "kwietnia": 4,
    "apr": 4,
    "april": 4,
    "maj": 5,
    "maja": 5,
    "may": 5,
    "cze": 6,
    "czerwiec": 6,
    "czerwca": 6,
    "jun": 6,
    "june": 6,
    "lip": 7,
    "lipiec": 7,
    "lipca": 7,
    "jul": 7,
    "july": 7,
    "sie": 8,
    "sierpień": 8,
    "sierpnia": 8,
    "aug": 8,
    "august": 8,
    "wrz": 9,
    "wrzesień": 9,
    "września": 9,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "paź": 10,
    "październik": 10,
    "października": 10,
    "oct": 10,
    "october": 10,
    "lis": 11,
    "listopad": 11,
    "listopada": 11,
    "nov": 11,
    "november": 11,
    "gru": 12,
    "grudzień": 12,
    "grudnia": 12,
    "dec": 12,
    "december": 12,
}
ROMAN_MONTH_TO_NUMBER = {
    "i": 1,
    "ii": 2,
    "iii": 3,
    "iv": 4,
    "v": 5,
    "vi": 6,
    "vii": 7,
    "viii": 8,
    "ix": 9,
    "x": 10,
    "xi": 11,
    "xii": 12,
}
MONTH_NAME_REGEX = "|".join(
    sorted((re.escape(name) for name in MONTH_NAME_TO_NUMBER), key=len, reverse=True)
)
ROMAN_MONTH_REGEX = "|".join(
    sorted((re.escape(name) for name in ROMAN_MONTH_TO_NUMBER), key=len, reverse=True)
)
TEXTUAL_DATE_PATTERN = re.compile(
    rf"""
    (?<!\w)
    (?:
        \d{{1,2}}(?:st|nd|rd|th|-go)?
        (?:\s+|\s*[-./]\s*)
        (?:{MONTH_NAME_REGEX})\.?
        (?:
            (?:\s+|\s*[-./,]\s*)
            \d{{2,4}}(?:\s*(?:r\.?|roku))?
        )?
        |
        (?:{MONTH_NAME_REGEX})\.?
        (?:\s+|\s*[-./]\s*)
        \d{{1,2}}(?:st|nd|rd|th)?
        (?:
            (?:\s+|\s*[-./,]\s*)
            \d{{2,4}}(?:\s*(?:r\.?|roku))?
        )?
        |
        (?:{MONTH_NAME_REGEX})\.?
        (?:\s+|\s*[-./]\s*)
        \d{{4}}(?:\s*(?:r\.?|roku))?
        |
        \d{{4}}
        (?:\s+|\s*[-./]\s*)
        (?:{MONTH_NAME_REGEX})\.?
        (?:\s+|\s*[-./]\s*)
        \d{{1,2}}(?:st|nd|rd|th)?
        |
        \d{{4}}
        (?:\s+|\s*[-./]\s*)
        (?:{MONTH_NAME_REGEX})\.?
        |
        \d{{1,2}}
        (?:\s+|\s*[-./]\s*)
        (?:{ROMAN_MONTH_REGEX})
        (?:\s+|\s*[-./]\s*)
        \d{{2,4}}(?:\s*(?:r\.?|roku))?
    )
    (?!\w)
    """,
    re.IGNORECASE | re.VERBOSE,
)
DATE_TIME_PATTERN = re.compile(
    r"""
    (?<!\w)
    (?:
        \d{1,4}(?P<datetime_sep>[-./_\\])\d{1,4}
        (?:(?P=datetime_sep)\d{1,4})?
        |
        \d{6,8}
    )
    [Tt\s]
    (?:
        (?:[01]?\d|2[0-3]):[0-5]\d
        (?::[0-5]\d(?:[.,]\d{1,6})?)?
        |
        (?:[01]\d|2[0-3])[0-5]\d(?:[0-5]\d)?
    )
    (?:Z|[+-](?:[01]\d|2[0-3]):?[0-5]\d)?
    (?!\w)
    """,
    re.IGNORECASE | re.VERBOSE,
)
ISO_WEEK_DATE_PATTERN = re.compile(
    r"(?<!\w)\d{4}-?W\d{2}(?:-?[1-7])?(?!\w)", re.IGNORECASE
)
ISO_ORDINAL_DATE_PATTERN = re.compile(r"(?<!\w)\d{4}-?\d{3}(?!\w)")
SPACE_DATE_PATTERN = re.compile(
    r"(?<!\w)\d{1,4}\s+\d{1,2}\s+\d{1,4}(?!\w)"
)
DELIMITED_DATE_PATTERN = re.compile(
    r"(?<!\w)\d{1,4}(?P<date_sep>[-./_\\])\d{1,4}"
    r"(?:(?P=date_sep)\d{1,4})?(?!\w)"
)
COMPACT_DATE_PATTERN = re.compile(r"(?<!\d)\d{6,8}(?!\d)")
STATIC_OUTPUT_LABELS = {
    "Opis Zgłoszenia:",
    "Komentarz dla użytkownika:",
    "Komentarz techniczny:",
    "[OPIS]",
    "[SQL]",
}
PERSON_FALSE_POSITIVE_WORDS = {
    "adres",
    "aplikacja",
    "błąd",
    "cnk",
    "dane",
    "data",
    "email",
    "komentarz",
    "konto",
    "mail",
    "microsoft",
    "numer",
    "odpowiedź",
    "opis",
    "problem",
    "sql",
    "system",
    "techniczny",
    "telefon",
    "tytuł",
    "użytkownika",
    "windows",
    "word",
    "zgłoszenia",
}
THEME_PALETTES = {
    "light": {
        "accent": "#0f766e",
        "accent_active": "#0b5f59",
        "accent_fg": "#ffffff",
        "accent_soft": "#d8f3ef",
        "bg": "#f1f4f6",
        "border": "#d9e0e5",
        "button": "#f7f9fa",
        "button_active": "#e9eef1",
        "error": "#b42318",
        "input_bg": "#fbfcfd",
        "input_fg": "#172026",
        "muted": "#66727d",
        "panel": "#ffffff",
        "scroll_thumb": "#b8c3ca",
        "scroll_track": "#edf1f3",
        "select_bg": "#bfe8e3",
        "select_fg": "#102a28",
        "success": "#087a55",
        "surface": "#f5f8f9",
        "text": "#172026",
    },
    "dark": {
        "accent": "#39b9ac",
        "accent_active": "#4ccbbd",
        "accent_fg": "#071a18",
        "accent_soft": "#173d39",
        "bg": "#101315",
        "border": "#2b3339",
        "button": "#22282d",
        "button_active": "#2c343a",
        "error": "#ff7a8a",
        "input_bg": "#14191d",
        "input_fg": "#eaf0f3",
        "muted": "#96a2aa",
        "panel": "#191e22",
        "scroll_thumb": "#3a464e",
        "scroll_track": "#151a1e",
        "select_bg": "#185c55",
        "select_fg": "#f4fbfa",
        "success": "#45d49a",
        "surface": "#20262b",
        "text": "#eef3f5",
    },
}
THEME_MODE_TO_KEY = {"Jasny": "light", "Ciemny": "dark"}
THEME_KEY_TO_MODE = {"light": "Jasny", "dark": "Ciemny"}


def px_font(size_px, family="Segoe UI"):
    # Tkinter interprets negative font sizes as pixels.
    return (family, -size_px)


class ScribblesoApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Scribbleso")
        self.geometry("1180x760")
        self.minsize(980, 660)

        self.settings = self._load_settings()
        self.theme_mode = tk.StringVar(value=self._settings_theme_mode())
        self.palette = THEME_PALETTES[self._theme_key()]
        self.anonymize_enabled = tk.BooleanVar(value=True)
        self.sql_enabled = tk.BooleanVar(value=False)
        self.cnk_enabled = tk.BooleanVar(value=False)
        self.cnk_count = tk.IntVar(value=1)
        self.cnk_items = []
        self._word_com_initialized = False
        self.current_date = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d"))
        self.output_file_path = tk.StringVar(value=self._settings_output_file_path())
        self.file_status = tk.StringVar(value="")

        self._configure_style()
        self._build_layout()
        self._bind_updates()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self._update_preview()

    def _load_settings(self):
        try:
            settings = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return {}
        return settings if isinstance(settings, dict) else {}

    def _settings_theme_mode(self):
        theme = self.settings.get("theme")
        if theme in THEME_KEY_TO_MODE:
            return THEME_KEY_TO_MODE[theme]
        if theme in THEME_MODE_TO_KEY:
            return theme
        return "Jasny"

    def _settings_output_file_path(self):
        path = self.settings.get("output_file_path", "")
        return path if isinstance(path, str) else ""

    def _theme_key(self):
        return THEME_MODE_TO_KEY.get(self.theme_mode.get(), "light")

    def _save_settings(self):
        settings = {
            "theme": self._theme_key(),
            "output_file_path": self.output_file_path.get().strip(),
        }
        SETTINGS_FILE.parent.mkdir(parents=True, exist_ok=True)
        SETTINGS_FILE.write_text(
            json.dumps(settings, indent=2, ensure_ascii=True),
            encoding="utf-8",
        )

    def _on_close(self):
        try:
            self._save_settings()
        except OSError:
            pass
        finally:
            self._uninitialize_word_com()
            self.destroy()

    def _uninitialize_word_com(self):
        if not self.__dict__.get("_word_com_initialized", False):
            return

        try:
            import pythoncom

            pythoncom.CoUninitialize()
        finally:
            self._word_com_initialized = False

    def _configure_style(self):
        palette = self.palette
        self.configure(bg=palette["bg"])

        style = ttk.Style(self)
        self.style = style
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        style.configure("TFrame", background=palette["bg"])
        style.configure("Panel.TFrame", background=palette["panel"], relief="flat")
        style.configure(
            "Card.TFrame",
            background=palette["panel"],
            bordercolor=palette["border"],
            lightcolor=palette["border"],
            darkcolor=palette["border"],
            borderwidth=1,
            relief="solid",
        )
        style.configure("Toolbar.TFrame", background=palette["surface"], relief="flat")
        style.configure(
            "TLabel",
            background=palette["bg"],
            foreground=palette["text"],
            font=px_font(UI_FONT_SIZE),
        )
        style.configure(
            "Panel.TLabel",
            background=palette["panel"],
            foreground=palette["text"],
            font=px_font(UI_FONT_SIZE),
        )
        style.configure(
            "Field.Panel.TLabel",
            background=palette["panel"],
            foreground=palette["text"],
            font=("Segoe UI", -UI_SMALL_FONT_SIZE, "bold"),
        )
        style.configure(
            "Toolbar.TLabel",
            background=palette["surface"],
            foreground=palette["muted"],
            font=px_font(UI_SMALL_FONT_SIZE),
        )
        style.configure(
            "Heading.TLabel",
            background=palette["panel"],
            foreground=palette["text"],
            font=("Segoe UI", -18, "bold"),
        )
        style.configure(
            "Muted.Panel.TLabel",
            background=palette["panel"],
            foreground=palette["muted"],
            font=px_font(UI_SMALL_FONT_SIZE),
        )
        style.configure(
            "TButton",
            background=palette["button"],
            bordercolor=palette["border"],
            lightcolor=palette["border"],
            darkcolor=palette["border"],
            borderwidth=1,
            relief="solid",
            focusthickness=1,
            focuscolor=palette["accent"],
            foreground=palette["text"],
            font=px_font(UI_FONT_SIZE),
            padding=(13, 8),
        )
        style.map(
            "TButton",
            background=[
                ("disabled", palette["button"]),
                ("active", palette["button_active"]),
            ],
            foreground=[("disabled", palette["muted"])],
            bordercolor=[("focus", palette["accent"])],
        )
        style.configure(
            "Accent.TButton",
            background=palette["accent"],
            bordercolor=palette["accent"],
            lightcolor=palette["accent"],
            darkcolor=palette["accent"],
            borderwidth=1,
            relief="solid",
            foreground=palette["accent_fg"],
            font=px_font(UI_FONT_SIZE),
            padding=(14, 8),
        )
        style.map(
            "Accent.TButton",
            background=[
                ("disabled", palette["button"]),
                ("active", palette["accent_active"]),
            ],
            foreground=[("disabled", palette["muted"])],
            bordercolor=[("active", palette["accent_active"])],
        )
        style.configure(
            "Ghost.TButton",
            background=palette["panel"],
            bordercolor=palette["panel"],
            lightcolor=palette["panel"],
            darkcolor=palette["panel"],
            borderwidth=1,
            relief="flat",
            foreground=palette["muted"],
            font=px_font(UI_FONT_SIZE),
            padding=(10, 7),
        )
        style.map(
            "Ghost.TButton",
            background=[("active", palette["surface"])],
            foreground=[("active", palette["text"])],
            bordercolor=[("focus", palette["border"])],
        )
        style.configure(
            "TCheckbutton",
            background=palette["panel"],
            foreground=palette["text"],
            font=px_font(UI_FONT_SIZE),
            padding=(0, 3),
        )
        style.map(
            "TCheckbutton",
            background=[("active", palette["panel"])],
            foreground=[("disabled", palette["muted"])],
            indicatorcolor=[
                ("selected", palette["accent"]),
                ("!selected", palette["input_bg"]),
            ],
        )
        style.layout(
            "Toolbar.TCheckbutton",
            [
                (
                    "Checkbutton.padding",
                    {
                        "sticky": "nswe",
                        "children": [("Checkbutton.label", {"sticky": "nswe"})],
                    },
                )
            ],
        )
        style.configure(
            "Toolbar.TCheckbutton",
            background=palette["surface"],
            foreground=palette["muted"],
            bordercolor=palette["surface"],
            lightcolor=palette["surface"],
            darkcolor=palette["surface"],
            borderwidth=1,
            relief="solid",
            anchor="center",
            font=px_font(UI_SMALL_FONT_SIZE),
            padding=(10, 6),
        )
        style.map(
            "Toolbar.TCheckbutton",
            background=[
                ("selected", palette["accent_soft"]),
                ("active", palette["button_active"]),
            ],
            foreground=[
                ("selected", palette["accent"]),
                ("active", palette["text"]),
                ("disabled", palette["muted"]),
            ],
            bordercolor=[("selected", palette["accent"])],
        )
        style.configure(
            "TEntry",
            bordercolor=palette["border"],
            lightcolor=palette["border"],
            darkcolor=palette["border"],
            borderwidth=1,
            relief="solid",
            fieldbackground=palette["input_bg"],
            foreground=palette["input_fg"],
            insertcolor=palette["text"],
            font=px_font(UI_FONT_SIZE),
            padding=(10, 7),
        )
        style.map(
            "TEntry",
            bordercolor=[("focus", palette["accent"])],
            fieldbackground=[("readonly", palette["surface"])],
            foreground=[("readonly", palette["muted"])],
        )
        style.configure(
            "TSpinbox",
            arrowsize=12,
            bordercolor=palette["border"],
            fieldbackground=palette["input_bg"],
            foreground=palette["input_fg"],
            insertcolor=palette["text"],
            font=px_font(UI_FONT_SIZE),
            padding=(7, 6),
        )
        style.configure(
            "TCombobox",
            arrowcolor=palette["text"],
            bordercolor=palette["border"],
            fieldbackground=palette["input_bg"],
            foreground=palette["input_fg"],
            font=px_font(UI_FONT_SIZE),
            padding=(9, 6),
        )
        style.map(
            "TCombobox",
            fieldbackground=[("readonly", palette["input_bg"])],
            foreground=[("readonly", palette["input_fg"])],
            selectbackground=[("readonly", palette["input_bg"])],
            selectforeground=[("readonly", palette["input_fg"])],
        )
        style.layout(
            "Segment.TRadiobutton",
            [
                (
                    "Radiobutton.padding",
                    {
                        "sticky": "nswe",
                        "children": [("Radiobutton.label", {"sticky": "nswe"})],
                    },
                )
            ],
        )
        style.configure(
            "Segment.TRadiobutton",
            anchor="center",
            background=palette["surface"],
            foreground=palette["muted"],
            bordercolor=palette["border"],
            lightcolor=palette["border"],
            darkcolor=palette["border"],
            borderwidth=1,
            relief="solid",
            font=px_font(UI_SMALL_FONT_SIZE),
            padding=(10, 6),
        )
        style.map(
            "Segment.TRadiobutton",
            background=[
                ("selected", palette["accent_soft"]),
                ("active", palette["button_active"]),
            ],
            foreground=[
                ("selected", palette["accent"]),
                ("active", palette["text"]),
            ],
            bordercolor=[("selected", palette["accent"])],
        )
        style.configure(
            "Vertical.TScrollbar",
            background=palette["scroll_thumb"],
            troughcolor=palette["scroll_track"],
            bordercolor=palette["scroll_track"],
            lightcolor=palette["scroll_thumb"],
            darkcolor=palette["scroll_thumb"],
            arrowcolor=palette["muted"],
            borderwidth=0,
            relief="flat",
            arrowsize=10,
            width=12,
        )
        style.map(
            "Vertical.TScrollbar",
            background=[("active", palette["accent"])],
            arrowcolor=[("active", palette["text"])],
        )
        style.configure("TSeparator", background=palette["border"])
        style.configure(
            "Status.TLabel",
            background=palette["panel"],
            foreground=palette["muted"],
            font=px_font(UI_SMALL_FONT_SIZE),
        )

    def _change_theme(self, _event=None):
        self.palette = THEME_PALETTES[self._theme_key()]
        self._configure_style()
        self._apply_theme_to_widgets()

    def _apply_theme_to_widgets(self):
        self._apply_theme_to_widget_tree(self)
        self.preview_text.tag_configure("default", font=px_font(UI_FONT_SIZE))
        self.preview_text.tag_configure(
            "title", font=("Segoe UI", -UI_FONT_SIZE, "bold")
        )
        self.preview_text.tag_configure("description", font=px_font(UI_SMALL_FONT_SIZE))

    def _apply_theme_to_widget_tree(self, widget):
        palette = self.palette

        if isinstance(widget, tk.Canvas):
            widget.configure(
                bg=palette["panel"],
                highlightbackground=palette["border"],
                highlightcolor=palette["accent"],
            )
        elif isinstance(widget, tk.Text):
            widget.configure(
                bg=palette["input_bg"],
                fg=palette["input_fg"],
                highlightbackground=palette["border"],
                highlightcolor=palette["accent"],
                insertbackground=palette["text"],
                selectbackground=palette["select_bg"],
                selectforeground=palette["select_fg"],
            )

        for child in widget.winfo_children():
            self._apply_theme_to_widget_tree(child)

    def _build_layout(self):
        root = ttk.Frame(self, padding=20)
        root.grid(row=0, column=0, sticky="nsew")

        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        root.columnconfigure(0, weight=1, uniform="columns")
        root.columnconfigure(1, weight=1, uniform="columns")
        root.rowconfigure(0, weight=1)

        form_shell = ttk.Frame(root, style="Card.TFrame", padding=1)
        form_shell.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        form_shell.columnconfigure(0, weight=1)
        form_shell.rowconfigure(0, weight=1)

        self.form_canvas = tk.Canvas(
            form_shell,
            bg=self.palette["panel"],
            highlightthickness=0,
        )
        self.form_canvas.grid(row=0, column=0, sticky="nsew")

        form_scrollbar = ttk.Scrollbar(form_shell, orient="vertical", command=self.form_canvas.yview)
        form_scrollbar.grid(row=0, column=1, sticky="ns")
        self.form_canvas.configure(yscrollcommand=form_scrollbar.set)

        form_panel = ttk.Frame(self.form_canvas, style="Panel.TFrame", padding=22)
        self.form_canvas_window = self.form_canvas.create_window(
            (0, 0), window=form_panel, anchor="nw"
        )
        form_panel.bind("<Configure>", self._update_form_scrollregion)
        self.form_canvas.bind("<Configure>", self._resize_form_panel)
        self._bind_form_mousewheel()
        form_panel.columnconfigure(0, weight=1)

        form_header = ttk.Frame(form_panel, style="Panel.TFrame")
        form_header.grid(row=0, column=0, sticky="ew", pady=(0, 16))
        form_header.columnconfigure(0, weight=1)

        ttk.Label(form_header, text="Dane zgłoszenia", style="Heading.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Button(
            form_header,
            text="Wyczyść",
            command=self.clear_form,
            style="Ghost.TButton",
        ).grid(
            row=0, column=1, sticky="e"
        )

        options_frame = ttk.Frame(
            form_panel, style="Toolbar.TFrame", padding=(12, 9)
        )
        options_frame.grid(row=1, column=0, sticky="ew", pady=(0, 18))

        self.anonymize_checkbox = ttk.Checkbutton(
            options_frame,
            text="Anonimizuj",
            variable=self.anonymize_enabled,
            command=self._update_preview,
            style="Toolbar.TCheckbutton",
        )
        self.anonymize_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 18))

        self.sql_checkbox = ttk.Checkbutton(
            options_frame,
            text="SQL",
            variable=self.sql_enabled,
            command=self._toggle_sql_input,
            style="Toolbar.TCheckbutton",
        )
        self.sql_checkbox.grid(row=0, column=1, sticky="w", padx=(0, 18))

        self.cnk_checkbox = ttk.Checkbutton(
            options_frame,
            text="CNK",
            variable=self.cnk_enabled,
            command=self._toggle_cnk_input,
            style="Toolbar.TCheckbutton",
        )
        self.cnk_checkbox.grid(row=0, column=2, sticky="w")

        self.cnk_controls = ttk.Frame(options_frame, style="Toolbar.TFrame")
        self.cnk_controls.grid(row=1, column=0, columnspan=3, sticky="ew", pady=(10, 0))
        self.cnk_controls.columnconfigure(2, weight=1)

        ttk.Label(self.cnk_controls, text="Liczba par CNK", style="Toolbar.TLabel").grid(
            row=0, column=0, sticky="w", padx=(0, 8)
        )
        self.cnk_count_spinbox = ttk.Spinbox(
            self.cnk_controls,
            from_=1,
            to=20,
            width=5,
            textvariable=self.cnk_count,
            command=self._sync_cnk_fields,
        )
        self.cnk_count_spinbox.grid(row=0, column=1, sticky="w")
        self.cnk_controls.grid_remove()

        title_frame = ttk.Frame(form_panel, style="Panel.TFrame")
        title_frame.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        title_frame.columnconfigure(0, weight=1)

        ttk.Label(title_frame, text="Tytuł", style="Field.Panel.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Label(title_frame, text="Data", style="Field.Panel.TLabel").grid(
            row=0, column=1, sticky="w", padx=(12, 0)
        )

        self.title_entry = ttk.Entry(title_frame, font=px_font(UI_FONT_SIZE))
        self.title_entry.grid(row=1, column=0, sticky="ew", pady=(5, 0))
        self.date_entry = ttk.Entry(
            title_frame,
            width=12,
            textvariable=self.current_date,
            font=px_font(UI_FONT_SIZE),
            state="readonly",
        )
        self.date_entry.grid(row=1, column=1, sticky="e", padx=(12, 0), pady=(5, 0))

        ttk.Label(form_panel, text="Opis zgłoszenia", style="Field.Panel.TLabel").grid(
            row=3, column=0, sticky="nw"
        )
        self.description_text = self._make_text(
            form_panel, height=5, font_size=UI_SMALL_FONT_SIZE
        )
        self.description_text.grid(row=4, column=0, sticky="nsew", pady=(6, 14))

        ttk.Label(
            form_panel,
            text="Komentarz dla użytkownika",
            style="Field.Panel.TLabel",
        ).grid(
            row=5, column=0, sticky="nw"
        )
        self.user_comment_text = self._make_text(form_panel, height=5)
        self.user_comment_text.grid(row=6, column=0, sticky="nsew", pady=(6, 14))

        ttk.Label(
            form_panel,
            text="Komentarz techniczny",
            style="Field.Panel.TLabel",
        ).grid(
            row=7, column=0, sticky="nw"
        )
        self.technical_comment_text = self._make_text(form_panel, height=5)
        self.technical_comment_text.grid(row=8, column=0, sticky="nsew", pady=(6, 12))

        self.sql_text = self._make_text(form_panel, height=4)

        self.cnk_frame = ttk.Frame(form_panel, style="Panel.TFrame")
        self.cnk_frame.columnconfigure(0, weight=1)

        self.cnk_fields_frame = ttk.Frame(self.cnk_frame, style="Panel.TFrame")
        self.cnk_fields_frame.grid(row=0, column=0, sticky="ew")
        self.cnk_fields_frame.columnconfigure(0, weight=1)

        preview_shell = ttk.Frame(root, style="Card.TFrame", padding=1)
        preview_shell.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        preview_shell.columnconfigure(0, weight=1)
        preview_shell.rowconfigure(0, weight=1)

        preview_panel = ttk.Frame(preview_shell, style="Panel.TFrame", padding=22)
        preview_panel.grid(row=0, column=0, sticky="nsew")
        preview_panel.columnconfigure(0, weight=1)
        preview_panel.rowconfigure(3, weight=1)

        header = ttk.Frame(preview_panel, style="Panel.TFrame")
        header.grid(row=0, column=0, sticky="ew", pady=(0, 16))
        header.columnconfigure(0, weight=1)

        ttk.Label(header, text="Wynik", style="Heading.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Label(header, text="Motyw", style="Muted.Panel.TLabel").grid(
            row=0, column=1, sticky="e", padx=(0, 8)
        )
        theme_switch = ttk.Frame(header, style="Toolbar.TFrame", padding=1)
        theme_switch.grid(row=0, column=2, sticky="e", padx=(0, 12))

        self.light_theme_button = ttk.Radiobutton(
            theme_switch,
            text="Jasny",
            variable=self.theme_mode,
            value="Jasny",
            command=self._change_theme,
            style="Segment.TRadiobutton",
        )
        self.light_theme_button.grid(row=0, column=0, sticky="nsew")
        self.dark_theme_button = ttk.Radiobutton(
            theme_switch,
            text="Ciemny",
            variable=self.theme_mode,
            value="Ciemny",
            command=self._change_theme,
            style="Segment.TRadiobutton",
        )
        self.dark_theme_button.grid(row=0, column=1, sticky="nsew", padx=(2, 0))

        self.copy_button = ttk.Button(
            header, text="Kopiuj", command=self.copy_result, style="Accent.TButton"
        )
        self.copy_button.grid(row=0, column=3, sticky="e")

        ttk.Separator(preview_panel, orient="horizontal").grid(
            row=1, column=0, sticky="ew", pady=(0, 16)
        )

        file_panel = ttk.Frame(preview_panel, style="Panel.TFrame")
        file_panel.grid(row=2, column=0, sticky="ew", pady=(0, 16))
        file_panel.columnconfigure(0, weight=1)

        ttk.Label(file_panel, text="Plik docelowy", style="Field.Panel.TLabel").grid(
            row=0, column=0, columnspan=3, sticky="w", pady=(0, 6)
        )
        self.output_file_entry = ttk.Entry(
            file_panel,
            textvariable=self.output_file_path,
            font=px_font(UI_FONT_SIZE),
        )
        self.output_file_entry.grid(row=1, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(file_panel, text="Wybierz", command=self.choose_output_file).grid(
            row=1, column=1, sticky="e", padx=(0, 8)
        )
        self.append_file_button = ttk.Button(
            file_panel,
            text="Dopisz",
            command=self.append_result_to_file,
            style="Accent.TButton",
        )
        self.append_file_button.grid(row=1, column=2, sticky="e")

        self.file_status_label = ttk.Label(
            file_panel, textvariable=self.file_status, style="Status.TLabel"
        )
        self.file_status_label.grid(
            row=2, column=0, columnspan=3, sticky="w", pady=(6, 0)
        )

        preview_frame = ttk.Frame(preview_panel, style="Panel.TFrame")
        preview_frame.grid(row=3, column=0, sticky="nsew")
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)

        self.preview_text = tk.Text(
            preview_frame,
            wrap="word",
            font=px_font(UI_FONT_SIZE),
            bg=self.palette["input_bg"],
            fg=self.palette["input_fg"],
            relief="flat",
            borderwidth=0,
            highlightthickness=1,
            highlightbackground=self.palette["border"],
            highlightcolor=self.palette["accent"],
            insertbackground=self.palette["text"],
            padx=16,
            pady=14,
            selectbackground=self.palette["select_bg"],
            selectforeground=self.palette["select_fg"],
            undo=False,
        )
        self.preview_text.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(preview_frame, orient="vertical", command=self.preview_text.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.preview_text.configure(yscrollcommand=scrollbar.set)

        self.preview_text.tag_configure("default", font=px_font(UI_FONT_SIZE))
        self.preview_text.tag_configure(
            "title", font=("Segoe UI", -UI_FONT_SIZE, "bold")
        )
        self.preview_text.tag_configure("description", font=px_font(UI_SMALL_FONT_SIZE))
        self.preview_text.configure(state="disabled")

    def _make_text(self, parent, height, font_size=UI_FONT_SIZE):
        frame = ttk.Frame(parent, style="Panel.TFrame")
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(0, weight=1)

        text = tk.Text(
            frame,
            height=height,
            wrap="word",
            font=px_font(font_size),
            bg=self.palette["input_bg"],
            fg=self.palette["input_fg"],
            relief="flat",
            borderwidth=0,
            highlightthickness=1,
            highlightbackground=self.palette["border"],
            highlightcolor=self.palette["accent"],
            padx=10,
            pady=9,
            insertbackground=self.palette["text"],
            selectbackground=self.palette["select_bg"],
            selectforeground=self.palette["select_fg"],
        )
        text.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=text.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        text.configure(yscrollcommand=scrollbar.set)

        return frame

    def _update_form_scrollregion(self, _event=None):
        self.form_canvas.configure(scrollregion=self.form_canvas.bbox("all"))

    def _resize_form_panel(self, event):
        self.form_canvas.itemconfigure(self.form_canvas_window, width=event.width)

    def _bind_form_mousewheel(self):
        self.bind_all("<MouseWheel>", self._on_form_mousewheel, add="+")
        self.bind_all("<Button-4>", self._on_form_mousewheel, add="+")
        self.bind_all("<Button-5>", self._on_form_mousewheel, add="+")

    def _on_form_mousewheel(self, event):
        widget = self.winfo_containing(event.x_root, event.y_root)
        if not self._is_widget_inside(widget, self.form_canvas):
            return None

        if self._is_widget_inside(widget, self.preview_text) or isinstance(widget, tk.Text):
            return None

        units = self._mousewheel_units(event)
        if units:
            self.form_canvas.yview_scroll(units, "units")
            return "break"
        return None

    def _is_widget_inside(self, widget, parent):
        while widget is not None:
            if widget == parent:
                return True
            widget = widget.master
        return False

    def _mousewheel_units(self, event):
        if getattr(event, "num", None) == 4:
            return -3
        if getattr(event, "num", None) == 5:
            return 3

        delta = getattr(event, "delta", 0)
        if not delta:
            return 0

        steps = int(delta / 120)
        if steps == 0:
            steps = 1 if delta > 0 else -1
        return -steps * 3

    def _bind_updates(self):
        self.title_entry.bind("<KeyRelease>", lambda _event: self._update_preview())
        self.cnk_count_spinbox.bind(
            "<KeyRelease>", lambda _event: self.after_idle(self._sync_cnk_fields)
        )
        self.cnk_count_spinbox.bind("<FocusOut>", lambda _event: self._sync_cnk_fields())
        self.cnk_count_spinbox.bind("<Return>", lambda _event: self._sync_cnk_fields())

        for widget in self._input_text_widgets():
            self._bind_text_updates(widget)

    def _bind_text_updates(self, widget):
        widget.bind("<KeyRelease>", lambda _event: self._update_preview())
        widget.bind("<<Paste>>", lambda _event: self.after_idle(self._update_preview))
        widget.bind("<<Cut>>", lambda _event: self.after_idle(self._update_preview))

    def _input_text_widgets(self):
        return [
            self._text_from_frame(self.description_text),
            self._text_from_frame(self.user_comment_text),
            self._text_from_frame(self.technical_comment_text),
            self._text_from_frame(self.sql_text),
        ]

    def _text_from_frame(self, frame):
        for child in frame.winfo_children():
            if isinstance(child, tk.Text):
                return child
        raise RuntimeError("Nie znaleziono pola tekstowego.")

    def clear_form(self):
        self.title_entry.delete(0, "end")
        for frame in (
            self.description_text,
            self.user_comment_text,
            self.technical_comment_text,
            self.sql_text,
        ):
            self._clear_text_frame(frame)

        self.sql_enabled.set(False)
        self.sql_text.grid_remove()

        self.cnk_enabled.set(False)
        self.cnk_count.set(1)
        self._sync_cnk_fields()
        for item in self.cnk_items:
            self._clear_text_frame(item["comment_text"])
            self._clear_text_frame(item["response_text"])
        self.cnk_controls.grid_remove()
        self.cnk_frame.grid_remove()

        self.file_status.set("")
        self._reset_copy_button()
        self.after_idle(self._update_form_scrollregion)
        self._update_preview()

    def _clear_text_frame(self, frame):
        self._text_from_frame(frame).delete("1.0", "end")

    def _toggle_sql_input(self):
        if self.sql_enabled.get():
            self.sql_text.grid(row=9, column=0, sticky="nsew", pady=(0, 12))
        else:
            self.sql_text.grid_remove()
        self.after_idle(self._update_form_scrollregion)
        self._update_preview()

    def _toggle_cnk_input(self):
        if self.cnk_enabled.get():
            self.cnk_controls.grid()
            self.cnk_frame.grid(row=10, column=0, sticky="ew", pady=(0, 12))
            self._sync_cnk_fields()
            return
        else:
            self.cnk_controls.grid_remove()
            self.cnk_frame.grid_remove()
            self.after_idle(self._update_form_scrollregion)
        self._update_preview()

    def _sync_cnk_fields(self):
        count = self._get_cnk_count()

        while len(self.cnk_items) < count:
            self.cnk_items.append(self._make_cnk_item(len(self.cnk_items) + 1))

        while len(self.cnk_items) > count:
            item = self.cnk_items.pop()
            item["frame"].destroy()

        self._reindex_cnk_items()
        self.after_idle(self._update_form_scrollregion)
        self._update_preview()

    def _get_cnk_count(self):
        try:
            current_count = self.cnk_count.get()
            invalid_count = False
        except tk.TclError:
            current_count = 1
            invalid_count = True

        count = current_count
        count = max(1, min(20, count))
        if invalid_count or current_count != count:
            self.cnk_count.set(count)
        return count

    def _make_cnk_item(self, index):
        item_frame = ttk.Frame(self.cnk_fields_frame, style="Panel.TFrame")
        item_frame.columnconfigure(0, weight=1)

        comment_label = ttk.Label(
            item_frame, text=f"Komentarz na CNK {index}", style="Field.Panel.TLabel"
        )
        comment_label.grid(row=0, column=0, sticky="nw")
        comment_text = self._make_text(item_frame, height=4)
        comment_text.grid(row=1, column=0, sticky="nsew", pady=(5, 10))

        response_label = ttk.Label(
            item_frame,
            text=f"Odpowiedź użytkownika {index}",
            style="Field.Panel.TLabel",
        )
        response_label.grid(row=2, column=0, sticky="nw")
        response_text = self._make_text(item_frame, height=4)
        response_text.grid(row=3, column=0, sticky="nsew", pady=(5, 12))

        self._bind_text_updates(self._text_from_frame(comment_text))
        self._bind_text_updates(self._text_from_frame(response_text))

        return {
            "frame": item_frame,
            "comment_label": comment_label,
            "comment_text": comment_text,
            "response_label": response_label,
            "response_text": response_text,
        }

    def _reindex_cnk_items(self):
        for index, item in enumerate(self.cnk_items, start=1):
            item["frame"].grid(row=index - 1, column=0, sticky="ew")
            item["comment_label"].configure(text=f"Komentarz na CNK {index}")
            item["response_label"].configure(text=f"Odpowiedź użytkownika {index}")

    def _get_text(self, frame):
        return self._text_from_frame(frame).get("1.0", "end-1c").strip()

    def _append_block(self, sections, lines):
        if sections:
            sections.append(("default", ""))
        sections.extend(lines)

    def _build_sections(self):
        title = self.title_entry.get().strip()
        description = self._get_text(self.description_text)
        user_comment = self._get_text(self.user_comment_text)
        technical_comment = self._get_text(self.technical_comment_text)
        sql = self._get_text(self.sql_text) if self.sql_enabled.get() else ""

        sections = []
        if title:
            sections.append(("title", f"{title}    {self.current_date.get()}"))

        if description:
            self._append_block(
                sections,
                [
                    ("default", "Opis Zgłoszenia:"),
                    ("description", description),
                ],
            )

        if user_comment:
            self._append_block(
                sections,
                [
                    ("default", "Komentarz dla użytkownika:"),
                    ("default", user_comment),
                ],
            )

        if technical_comment:
            self._append_block(
                sections,
                [
                    ("default", "Komentarz techniczny:"),
                    ("default", "[OPIS]"),
                    ("default", technical_comment),
                ],
            )

        if sql:
            self._append_block(
                sections,
                [
                    ("default", "[SQL]"),
                    ("default", sql),
                ],
            )

        if self.cnk_enabled.get():
            for index, item in enumerate(self.cnk_items, start=1):
                cnk_comment = self._get_text(item["comment_text"])
                user_response = self._get_text(item["response_text"])

                if cnk_comment:
                    self._append_block(
                        sections,
                        [
                            ("default", f"Komentarz na CNK {index}:"),
                            ("default", cnk_comment),
                        ],
                    )

                if user_response:
                    self._append_block(
                        sections,
                        [
                            ("default", f"Odpowiedź Użytkownika {index}:"),
                            ("default", user_response),
                        ],
                    )

        if self.anonymize_enabled.get():
            return self._anonymize_sections(sections)

        return sections

    def _anonymize_sections(self, sections):
        replacements = {}
        counters = {"ADRES": 0, "EMAIL": 0, "ID": 0, "KLIENT": 0}
        protected_values = {}

        def replacement(match):
            kind = self._anonymized_match_kind(match)
            value = match.group(0)
            if kind == "KLIENT" and self._is_false_positive_person(value):
                return value

            key = (kind, value)

            if key not in replacements:
                counters[kind] += 1
                replacements[key] = f"[{kind}_{counters[kind]}]"

            return replacements[key]

        def anonymize_text(text):
            protected_text = self._protect_references(text, protected_values)
            anonymized_text = ANONYMIZE_PATTERN.sub(replacement, protected_text)
            return self._restore_references(anonymized_text, protected_values)

        return [
            (
                tag,
                anonymize_text(text)
                if self._should_anonymize_text(tag, text)
                else text,
            )
            for tag, text in sections
        ]

    def _protect_references(self, text, protected_values):
        protected_text = PROTECTED_REFERENCE_PATTERN.sub(
            lambda match: self._store_protected_value(match.group(0), protected_values),
            text,
        )
        date_patterns = (
            (TEXTUAL_DATE_PATTERN, self._is_textual_date),
            (DATE_TIME_PATTERN, self._is_date_time),
            (ISO_WEEK_DATE_PATTERN, self._is_iso_week_date),
            (ISO_ORDINAL_DATE_PATTERN, self._is_iso_ordinal_date),
            (SPACE_DATE_PATTERN, self._is_space_date),
            (DELIMITED_DATE_PATTERN, self._is_delimited_date),
            (COMPACT_DATE_PATTERN, self._is_compact_date),
        )
        for pattern, validator in date_patterns:
            protected_text = pattern.sub(
                lambda match, validator=validator: (
                    self._store_protected_value(match.group(0), protected_values)
                    if validator(match.group(0))
                    else match.group(0)
                ),
                protected_text,
            )
        return protected_text

    def _store_protected_value(self, value, protected_values):
        token = f"__SCRIBBLESO_KEEP_{len(protected_values)}__"
        protected_values[token] = value
        return token

    def _is_textual_date(self, value):
        normalized = value.casefold()
        normalized = re.sub(r"\b(?:roku|r)\.?\s*$", "", normalized).strip()
        normalized = re.sub(
            r"(?<=\d)(?:st|nd|rd|th|-go)\b", "", normalized
        )

        words = re.findall(r"[a-ząćęłńóśźż]+", normalized)
        month = next(
            (MONTH_NAME_TO_NUMBER[word.rstrip(".")] for word in words if word.rstrip(".") in MONTH_NAME_TO_NUMBER),
            None,
        )
        if month is None:
            month = next(
                (ROMAN_MONTH_TO_NUMBER[word] for word in words if word in ROMAN_MONTH_TO_NUMBER),
                None,
            )
        if month is None:
            return False

        number_tokens = re.findall(r"\d+", normalized)
        if len(number_tokens) == 1:
            token = number_tokens[0]
            if len(token) == 4:
                return self._is_valid_month_year(int(token), month)
            return self._is_valid_calendar_date(2000, month, int(token))

        if len(number_tokens) == 2:
            first_token, second_token = number_tokens
            if len(first_token) == 4 and len(second_token) <= 2:
                year_token, day_token = first_token, second_token
            elif len(second_token) in (2, 4) and len(first_token) <= 2:
                day_token, year_token = first_token, second_token
            else:
                return False
            return self._is_valid_calendar_date(
                self._normalized_year(year_token), month, int(day_token)
            )

        return False

    def _is_date_time(self, value):
        date_value = re.split(r"[Tt\s]", value, maxsplit=1)[0]
        if re.search(r"[-./_\\]", date_value):
            return self._is_delimited_date(date_value)
        return self._is_compact_date(date_value)

    def _is_iso_week_date(self, value):
        match = re.fullmatch(
            r"(\d{4})-?W(\d{2})(?:-?([1-7]))?", value, re.IGNORECASE
        )
        if match is None:
            return False

        try:
            datetime.fromisocalendar(
                int(match.group(1)), int(match.group(2)), int(match.group(3) or 1)
            )
        except ValueError:
            return False
        return True

    def _is_iso_ordinal_date(self, value):
        compact_value = value.replace("-", "")
        if len(compact_value) != 7:
            return False

        try:
            parsed = datetime.strptime(compact_value, "%Y%j")
        except ValueError:
            return False
        return parsed.year == int(compact_value[:4])

    def _is_space_date(self, value):
        normalized = re.sub(r"\s+", "-", value.strip())
        return self._is_delimited_date(normalized)

    def _is_delimited_date(self, value):
        parts = re.split(r"[-./_\\]", value)
        if len(parts) == 2:
            first, second = parts
            if len(first) == 4 and len(second) <= 2:
                return self._is_valid_month_year(int(first), int(second))
            if len(second) == 4 and len(first) <= 2:
                return self._is_valid_month_year(int(second), int(first))
            if len(first) <= 2 and len(second) <= 2:
                return any(
                    self._is_valid_calendar_date(2000, month, day)
                    for month, day in (
                        (int(first), int(second)),
                        (int(second), int(first)),
                    )
                )
            return False

        if len(parts) != 3:
            return False

        candidates = []
        for year_index, year_token in enumerate(parts):
            if len(year_token) not in (2, 4):
                continue
            remaining = [parts[index] for index in range(3) if index != year_index]
            if any(len(part) > 2 for part in remaining):
                continue
            year = self._normalized_year(year_token)
            candidates.extend(
                (
                    (year, int(remaining[0]), int(remaining[1])),
                    (year, int(remaining[1]), int(remaining[0])),
                )
            )

        return any(self._is_valid_calendar_date(*candidate) for candidate in candidates)

    def _is_compact_date(self, value):
        if len(value) == 7:
            return self._is_iso_ordinal_date(value)

        candidates = []
        if len(value) == 8:
            candidates.extend(
                (
                    (int(value[:4]), int(value[4:6]), int(value[6:8])),
                    (int(value[:4]), int(value[6:8]), int(value[4:6])),
                    (int(value[4:8]), int(value[2:4]), int(value[:2])),
                    (int(value[4:8]), int(value[:2]), int(value[2:4])),
                )
            )
        elif len(value) == 6:
            candidates.extend(
                (
                    (self._normalized_year(value[:2]), int(value[2:4]), int(value[4:6])),
                    (self._normalized_year(value[:2]), int(value[4:6]), int(value[2:4])),
                    (self._normalized_year(value[4:6]), int(value[2:4]), int(value[:2])),
                    (self._normalized_year(value[4:6]), int(value[:2]), int(value[2:4])),
                )
            )
        else:
            return False

        return any(self._is_valid_calendar_date(*candidate) for candidate in candidates)

    def _normalized_year(self, value):
        year = int(value)
        return 2000 + year if len(value) == 2 else year

    def _is_valid_month_year(self, year, month):
        return 1 <= year <= 9999 and 1 <= month <= 12

    def _is_valid_calendar_date(self, year, month, day):
        try:
            datetime(year, month, day)
        except ValueError:
            return False
        return True

    def _restore_references(self, text, protected_values):
        for token, value in protected_values.items():
            text = text.replace(token, value)
        return text

    def _anonymized_match_kind(self, match):
        if match.group("address"):
            return "ADRES"
        if match.group("email"):
            return "EMAIL"
        if match.group("person"):
            return "KLIENT"
        return "ID"

    def _should_anonymize_text(self, tag, text):
        if tag == "title":
            return False

        if not text or text in STATIC_OUTPUT_LABELS:
            return False

        return not re.fullmatch(r"(Komentarz na CNK|Odpowiedź Użytkownika) \d+:", text)

    def _is_false_positive_person(self, value):
        words = re.split(r"[\s-]+", value.casefold())
        return any(word in PERSON_FALSE_POSITIVE_WORDS for word in words)

    def _update_preview(self):
        sections = self._build_sections()

        self.preview_text.configure(state="normal")
        self.preview_text.delete("1.0", "end")

        for index, (tag, text) in enumerate(sections):
            if index > 0:
                self.preview_text.insert("end", "\n")
            self.preview_text.insert("end", text, tag)

        self.preview_text.configure(state="disabled")

    def copy_result(self):
        sections = self._build_sections()
        plain_text = self._plain_text(sections)
        html_text = self._html_text(sections)

        try:
            self._copy_html_to_clipboard(html_text, plain_text)
        except Exception:
            self.clipboard_clear()
            self.clipboard_append(plain_text)
            self.update_idletasks()

        self._show_copy_feedback()

    def _show_copy_feedback(self):
        self.copy_button.configure(text="Skopiowano", state="disabled")
        self.after(3000, self._reset_copy_button)

    def _reset_copy_button(self):
        self.copy_button.configure(text="Kopiuj", state="normal")

    def choose_output_file(self):
        path = filedialog.asksaveasfilename(
            title="Wybierz plik do dopisywania",
            defaultextension=".txt",
            filetypes=[
                ("Pliki tekstowe i Word", "*.txt *.docx"),
                ("Pliki tekstowe", "*.txt"),
                ("Dokumenty Word", "*.docx"),
                ("Wszystkie pliki", "*.*"),
            ],
            confirmoverwrite=False,
        )
        if path:
            self.output_file_path.set(path)
            self.file_status.set("")

    def append_result_to_file(self):
        sections = self._build_sections()
        plain_text = self._plain_text(sections).strip()
        if not plain_text:
            self._show_file_status("Brak wyniku do dopisania.", success=False)
            return

        path_text = self.output_file_path.get().strip()
        if not path_text:
            self._show_file_status("Najpierw wybierz plik .txt albo .docx.", success=False)
            return

        path = Path(path_text)
        suffix = path.suffix.lower()

        try:
            if suffix == ".txt":
                self._append_to_txt(path, plain_text)
            elif suffix == ".docx":
                self._append_to_docx(path, sections)
            else:
                self._show_file_status("Obsługiwane są tylko pliki .txt i .docx.", success=False)
                return
        except Exception as exc:
            self._show_file_status(f"Nie udało się dopisać: {exc}", success=False)
            return

        self._show_file_status("Dopisano wynik do pliku.", success=True)

    def _append_to_txt(self, path, plain_text):
        try:
            content, encoding = self._read_text_file(path)
        except OSError:
            self._append_to_open_txt(path, plain_text, "utf-8")
            return

        content = content.rstrip()
        new_content = plain_text if not content else f"{content}\n\n\n{plain_text}"
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            path.write_text(new_content, encoding=encoding)
        except OSError:
            self._append_to_open_txt(path, plain_text, encoding)

    def _append_to_open_txt(self, path, plain_text, encoding):
        separator = "" if not path.exists() or path.stat().st_size == 0 else "\n\n\n"
        with path.open("a", encoding=encoding, newline="") as file:
            file.write(f"{separator}{plain_text}")

    def _read_text_file(self, path):
        if not path.exists():
            return "", "utf-8"

        data = path.read_bytes()
        if data.startswith(b"\xef\xbb\xbf"):
            return data.decode("utf-8-sig"), "utf-8-sig"

        for encoding in ("utf-8", "cp1250", "cp1252"):
            try:
                return data.decode(encoding), encoding
            except UnicodeDecodeError:
                pass

        return data.decode("utf-8", errors="replace"), "utf-8"

    def _append_to_docx(self, path, sections):
        if self._append_to_open_word_document(path, sections):
            return

        try:
            self._append_to_docx_file(path, sections)
        except OSError as exc:
            if self._append_to_open_word_document(path, sections):
                return
            raise exc

    def _append_to_docx_file(self, path, sections):
        from docx import Document
        from docx.shared import Pt

        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document(path) if path.exists() else Document()

        if self._docx_has_text(document):
            self._trim_docx_trailing_empty_paragraphs(document)
            document.add_paragraph()
            document.add_paragraph()
        elif len(document.paragraphs) == 1 and not document.paragraphs[0].text.strip():
            self._remove_docx_paragraph(document.paragraphs[0])

        for tag, text in sections:
            size = DESCRIPTION_FONT_SIZE if tag == "description" else DEFAULT_FONT_SIZE
            if not text:
                document.add_paragraph()
                continue

            for line in text.split("\n"):
                paragraph = (
                    document.add_paragraph(style="Heading 1")
                    if tag == "title"
                    else document.add_paragraph()
                )
                run = paragraph.add_run(line)
                if tag != "title":
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(size)

        document.save(path)

    def _append_to_open_word_document(self, path, sections):
        try:
            import pythoncom
            import win32com.client
        except ImportError:
            return False

        target_document_found = False
        previous_display_alerts = None
        self._initialize_word_com(pythoncom)
        try:
            try:
                word = win32com.client.GetActiveObject("Word.Application")
            except Exception:
                return False

            document = self._find_open_word_document(word, path)
            if document is None:
                return False

            target_document_found = True
            previous_display_alerts = word.DisplayAlerts
            word.DisplayAlerts = 0
            self._append_sections_to_word_document(document, sections)
            document.Save()
            return True
        except Exception as exc:
            if target_document_found:
                raise RuntimeError(
                    f"Nie udało się dopisać do otwartego dokumentu Word: {exc}"
                ) from exc
            return False
        finally:
            if previous_display_alerts is not None:
                try:
                    word.DisplayAlerts = previous_display_alerts
                except Exception:
                    pass

    def _initialize_word_com(self, pythoncom):
        if self.__dict__.get("_word_com_initialized", False):
            return

        pythoncom.CoInitialize()
        self._word_com_initialized = True

    def _find_open_word_document(self, word, path):
        target_path = self._normalized_path(path)
        for document in word.Documents:
            try:
                if self._normalized_path(document.FullName) == target_path:
                    return document
            except Exception:
                pass
        return None

    def _normalized_path(self, path):
        return os.path.normcase(os.path.abspath(str(path)))

    def _append_sections_to_word_document(self, document, sections):
        if self._word_document_has_text(document):
            self._trim_word_trailing_empty_paragraphs(document)
            self._append_word_paragraph(document, "")
            self._append_word_paragraph(document, "")

        for tag, text in sections:
            size = DESCRIPTION_FONT_SIZE if tag == "description" else DEFAULT_FONT_SIZE
            if not text:
                self._append_word_paragraph(document, "")
                continue

            for line in text.split("\n"):
                self._append_word_paragraph(document, line, tag=tag, size=size)

    def _append_word_paragraph(self, document, text, tag="default", size=DEFAULT_FONT_SIZE):
        wd_style_normal = -1
        wd_style_heading_1 = -2
        insertion_point = max(0, document.Content.End - 1)
        insertion_range = document.Range(insertion_point, insertion_point)
        prefix = "\r" if self._word_document_has_text(document) else ""
        insertion_range.Text = f"{prefix}{text}"
        paragraph = document.Paragraphs(document.Paragraphs.Count)

        if tag == "title":
            paragraph.Style = wd_style_heading_1
        else:
            paragraph.Style = wd_style_normal
            paragraph.Range.Font.Name = "Segoe UI"
            paragraph.Range.Font.Size = size

    def _word_document_has_text(self, document):
        text = document.Content.Text.replace("\r", "").replace("\x07", "").strip()
        return bool(text)

    def _trim_word_trailing_empty_paragraphs(self, document):
        while document.Paragraphs.Count > 1:
            paragraph = document.Paragraphs(document.Paragraphs.Count)
            text = paragraph.Range.Text.replace("\r", "").replace("\x07", "").strip()
            if text:
                break
            paragraph.Range.Delete()

    def _docx_has_text(self, document):
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                return True

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        return True

        return False

    def _trim_docx_trailing_empty_paragraphs(self, document):
        body = document._body._element
        children = list(body)
        index = len(children) - 1

        if index >= 0 and children[index].tag.endswith("}sectPr"):
            index -= 1

        while index >= 0:
            element = children[index]
            if not element.tag.endswith("}p") or self._docx_paragraph_has_content(element):
                break

            body.remove(element)
            index -= 1

    def _docx_paragraph_has_content(self, paragraph_element):
        text = "".join(
            element.text or ""
            for element in paragraph_element.iter()
            if element.tag.endswith("}t")
        )
        if text.strip():
            return True

        content_tags = ("}drawing", "}pict", "}object", "}fldSimple")
        return any(
            element.tag.endswith(content_tags)
            for element in paragraph_element.iter()
        )

    def _remove_docx_paragraph(self, paragraph):
        element = paragraph._element
        element.getparent().remove(element)

    def _show_file_status(self, message, success):
        foreground = self.palette["success"] if success else self.palette["error"]
        self.file_status_label.configure(foreground=foreground)
        self.file_status.set(message)

    def _plain_text(self, sections):
        return "\n".join(text for _tag, text in sections)

    def _html_text(self, sections):
        lines = []
        for tag, text in sections:
            size = DESCRIPTION_FONT_SIZE if tag == "description" else DEFAULT_FONT_SIZE
            escaped = html.escape(text).replace("\n", "<br>") or "<br>"
            lines.append(
                f'<div style="font-family: Segoe UI, Arial, sans-serif; '
                f'font-size: {size}px; line-height: 1.35;">{escaped}</div>'
            )
        return "".join(lines)

    def _copy_html_to_clipboard(self, html_text, plain_text):
        if sys.platform.startswith("win"):
            self._copy_html_to_windows_clipboard(html_text, plain_text)
            return

        self.clipboard_clear()
        self.clipboard_append(plain_text)
        self.update_idletasks()

    def _copy_html_to_windows_clipboard(self, html_text, plain_text):
        import ctypes
        from ctypes import wintypes

        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32

        cf_unicode_text = 13
        cf_html = user32.RegisterClipboardFormatW("HTML Format")

        user32.OpenClipboard.argtypes = [wintypes.HWND]
        user32.OpenClipboard.restype = wintypes.BOOL
        user32.EmptyClipboard.argtypes = []
        user32.EmptyClipboard.restype = wintypes.BOOL
        user32.SetClipboardData.argtypes = [wintypes.UINT, wintypes.HANDLE]
        user32.SetClipboardData.restype = wintypes.HANDLE
        user32.CloseClipboard.argtypes = []
        user32.CloseClipboard.restype = wintypes.BOOL
        kernel32.GlobalAlloc.argtypes = [wintypes.UINT, ctypes.c_size_t]
        kernel32.GlobalAlloc.restype = wintypes.HGLOBAL
        kernel32.GlobalLock.argtypes = [wintypes.HGLOBAL]
        kernel32.GlobalLock.restype = wintypes.LPVOID
        kernel32.GlobalUnlock.argtypes = [wintypes.HGLOBAL]
        kernel32.GlobalUnlock.restype = wintypes.BOOL
        kernel32.GlobalFree.argtypes = [wintypes.HGLOBAL]
        kernel32.GlobalFree.restype = wintypes.HGLOBAL

        def set_clipboard_data(format_id, data):
            handle = kernel32.GlobalAlloc(0x0042, len(data))
            if not handle:
                raise OSError("GlobalAlloc nie przydzielił pamięci.")

            locked = kernel32.GlobalLock(handle)
            if not locked:
                kernel32.GlobalFree(handle)
                raise OSError("GlobalLock nie zablokował pamięci.")

            ctypes.memmove(locked, data, len(data))
            kernel32.GlobalUnlock(handle)

            if not user32.SetClipboardData(format_id, handle):
                kernel32.GlobalFree(handle)
                raise OSError("SetClipboardData nie zapisał danych.")

        fragment = html_text
        clipboard_html = self._make_cf_html(fragment)
        unicode_text = plain_text.encode("utf-16le") + b"\x00\x00"

        if not user32.OpenClipboard(None):
            raise OSError("Schowek jest aktualnie niedostępny.")

        try:
            if not user32.EmptyClipboard():
                raise OSError("Nie udało się wyczyścić schowka.")
            set_clipboard_data(cf_unicode_text, unicode_text)
            set_clipboard_data(cf_html, clipboard_html)
        finally:
            user32.CloseClipboard()

    def _make_cf_html(self, fragment):
        start_fragment = "<!--StartFragment-->"
        end_fragment = "<!--EndFragment-->"
        html_payload = (
            "<html><body>"
            + start_fragment
            + fragment
            + end_fragment
            + "</body></html>"
        )

        header_template = (
            "Version:0.9\r\n"
            "StartHTML:{start_html:010d}\r\n"
            "EndHTML:{end_html:010d}\r\n"
            "StartFragment:{start_fragment:010d}\r\n"
            "EndFragment:{end_fragment:010d}\r\n"
        )
        placeholder = header_template.format(
            start_html=0,
            end_html=0,
            start_fragment=0,
            end_fragment=0,
        )

        start_html_offset = len(placeholder.encode("utf-8"))
        start_index = html_payload.index(start_fragment) + len(start_fragment)
        end_index = html_payload.index(end_fragment)
        start_fragment_offset = start_html_offset + len(html_payload[:start_index].encode("utf-8"))
        end_fragment_offset = start_html_offset + len(html_payload[:end_index].encode("utf-8"))
        html_bytes = html_payload.encode("utf-8")
        end_html_offset = start_html_offset + len(html_bytes)

        header = header_template.format(
            start_html=start_html_offset,
            end_html=end_html_offset,
            start_fragment=start_fragment_offset,
            end_fragment=end_fragment_offset,
        )
        return header.encode("utf-8") + html_bytes + b"\x00"


if __name__ == "__main__":
    app = ScribblesoApp()
    app.mainloop()
